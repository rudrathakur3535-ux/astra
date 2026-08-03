"""
Document Loader Module for Project Astra.
Parses multi-format documents (PDF, DOCX, Markdown, TXT, HTML, Python/Source Code, and Git Repositories).
"""

import os
from typing import List, Tuple, Dict, Any
from app.models.document import Document, DocumentType
from app.utils.logger import logger


class DocumentLoader:
    """
    Unified loader for extracting raw text and structure from multiple document formats.
    """

    def load_document(self, filepath: str) -> Tuple[Document, List[Tuple[str, int, str]]]:
        """
        Loads a single document file.

        Returns:
            Tuple[Document, List[Tuple[text, page_number, section_heading]]]
        """
        filepath = os.path.abspath(filepath)
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"Document file not found: {filepath}")

        filename = os.path.basename(filepath)
        ext = os.path.splitext(filename)[1].lower()
        file_size = os.path.getsize(filepath)

        doc_type = self._detect_type(ext)
        doc = Document(
            filepath=filepath,
            doc_type=doc_type,
            title=filename,
            file_size_bytes=file_size
        )

        segments = self._parse_file(filepath, doc_type)
        logger.info(f"Loaded document '{filename}' ({doc_type.value}, {len(segments)} segments, {file_size} bytes).")
        return doc, segments

    def load_git_repository(
        self,
        repo_path: str,
        allowed_extensions: Tuple[str, ...] = (".py", ".md", ".txt", ".json", ".yaml", ".yml", ".sql")
    ) -> List[Tuple[Document, List[Tuple[str, int, str]]]]:
        """
        Recursively scans a Git repository directory and loads all source code and documentation files.
        """
        repo_path = os.path.abspath(repo_path)
        if not os.path.exists(repo_path):
            raise FileNotFoundError(f"Git repository path not found: {repo_path}")

        loaded_documents = []
        ignored_dirs = {".git", ".pytest_cache", "__pycache__", "venv", "node_modules", "dist", "build", "chroma_db", "chroma_knowledge"}

        for root, dirs, files in os.walk(repo_path):
            # Exclude ignored directories
            dirs[:] = [d for d in dirs if d not in ignored_dirs]

            for file in files:
                ext = os.path.splitext(file)[1].lower()
                if ext in allowed_extensions:
                    full_path = os.path.join(root, file)
                    try:
                        doc, segments = self.load_document(full_path)
                        doc.collection = "code_repos"
                        loaded_documents.append((doc, segments))
                    except Exception as e:
                        logger.warning(f"Skipping repository file '{file}' due to parse error: {e}")

        logger.info(f"Loaded {len(loaded_documents)} repository files from '{repo_path}'.")
        return loaded_documents

    def _detect_type(self, ext: str) -> DocumentType:
        if ext == ".pdf":
            return DocumentType.PDF
        elif ext == ".docx":
            return DocumentType.DOCX
        elif ext in (".md", ".markdown"):
            return DocumentType.MARKDOWN
        elif ext in (".html", ".htm"):
            return DocumentType.HTML
        elif ext in (".py", ".js", ".ts", ".java", ".c", ".cpp", ".sql", ".json"):
            return DocumentType.CODE
        else:
            return DocumentType.TXT

    def _parse_file(self, filepath: str, doc_type: DocumentType) -> List[Tuple[str, int, str]]:
        """Parses file content according to its format."""

        # 1. PDF Parsing via pypdf
        if doc_type == DocumentType.PDF:
            return self._parse_pdf(filepath)

        # 2. DOCX Parsing via python-docx
        elif doc_type == DocumentType.DOCX:
            return self._parse_docx(filepath)

        # 3. Text / Markdown / Code / HTML Parsing
        else:
            return self._parse_text_file(filepath, doc_type)

    def _parse_pdf(self, filepath: str) -> List[Tuple[str, int, str]]:
        segments = []
        try:
            import pypdf
            reader = pypdf.PdfReader(filepath)
            for idx, page in enumerate(reader.pages, 1):
                text = page.extract_text() or ""
                if text.strip():
                    segments.append((text.strip(), idx, f"Page {idx}"))
        except Exception as e:
            logger.warning(f"pypdf failed on '{filepath}' ({e}). Reading raw file fallback.")
            segments = self._parse_text_file(filepath, DocumentType.TXT)
        return segments

    def _parse_docx(self, filepath: str) -> List[Tuple[str, int, str]]:
        segments = []
        try:
            import docx
            doc = docx.Document(filepath)
            current_heading = "Document Content"
            paragraph_buffer = []

            for p in doc.paragraphs:
                text = p.text.strip()
                if not text:
                    continue

                if p.style.name.startswith("Heading"):
                    if paragraph_buffer:
                        segments.append(("\n".join(paragraph_buffer), 1, current_heading))
                        paragraph_buffer = []
                    current_heading = text
                else:
                    paragraph_buffer.append(text)

            if paragraph_buffer:
                segments.append(("\n".join(paragraph_buffer), 1, current_heading))
        except Exception as e:
            logger.warning(f"python-docx failed on '{filepath}' ({e}). Reading raw text fallback.")
            segments = self._parse_text_file(filepath, DocumentType.TXT)
        return segments

    def _parse_text_file(self, filepath: str, doc_type: DocumentType) -> List[Tuple[str, int, str]]:
        try:
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
        except Exception as e:
            logger.error(f"Error reading file '{filepath}': {e}")
            content = ""

        # Extract markdown headings if present
        if doc_type == DocumentType.MARKDOWN:
            lines = content.splitlines()
            segments = []
            current_heading = "Introduction"
            buffer = []

            for line in lines:
                if line.startswith("#"):
                    if buffer:
                        segments.append(("\n".join(buffer), 1, current_heading))
                        buffer = []
                    current_heading = line.lstrip("#").strip()
                else:
                    buffer.append(line)

            if buffer:
                segments.append(("\n".join(buffer), 1, current_heading))
            return segments if segments else [(content, 1, "Full Content")]

        return [(content, 1, "Full Content")]
